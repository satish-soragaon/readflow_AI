"""Extract text from DOCX files, including table cells."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_docx_text(file_path: Path) -> str:
    """
    Extract all readable text from a DOCX file.

    Walks the document body in order, yielding paragraphs AND table cells so
    that content in tables is no longer silently dropped.
    """
    doc = Document(file_path)
    parts: list[str] = []

    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                # Deduplicate merged cells (python-docx repeats them)
                seen: set[str] = set()
                for cell in row.cells:
                    cell_id = cell._tc.xml  # unique per physical cell
                    if cell_id in seen:
                        continue
                    seen.add(cell_id)
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)

    return "\n".join(parts)


def _iter_blocks(doc: Document):
    """
    Yield Paragraph and Table objects in document body order.

    python-docx's doc.paragraphs and doc.tables are flat lists that lose
    relative ordering; iterating doc.element.body preserves it.
    """
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)
